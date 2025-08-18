import os, io, json, base64, textwrap, uuid
from typing import Optional, Dict, Any, List
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
import httpx
from google.cloud import storage
from google.cloud import firestore
from reportlab.pdfgen import canvas
from docx import Document
from pypdf import PdfReader

# Load .env only if present (handy for local dev; harmless on Cloud Run)
try:
    load_dotenv()
except Exception:
    pass

# --- Core Config ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
MODEL_NAME = os.getenv("MODEL_NAME", "gpt-5")
PROJECT_ID = os.getenv("GCP_PROJECT_ID", "project-manthan-468609")
BUCKET_NAME = os.getenv("GCS_BUCKET", "manthan-assets")
NAMESPACE = os.getenv("FIRESTORE_NAMESPACE", "prod")
ENABLE_IMAGE_GEN = os.getenv("ENABLE_IMAGE_GEN", "false").lower() == "true"

# Allow single origin, comma-separated list, or wildcard "*"
_raw_allowed = os.getenv("ALLOWED_ORIGIN", "http://localhost:3000").strip()
if "," in _raw_allowed:
    ALLOWED_ORIGINS = [o.strip() for o in _raw_allowed.split(",") if o.strip()]
else:
    ALLOWED_ORIGINS = ["*"] if _raw_allowed == "*" else [_raw_allowed]

# --- App ---
app = FastAPI(title="Manthan Creator Suite API", version="0.3.0")

# FastAPI CORS behavior: if using wildcard, must NOT set allow_credentials=True
_allow_credentials = not (len(ALLOWED_ORIGINS) == 1 and ALLOWED_ORIGINS[0] == "*")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=_allow_credentials,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Clients (lazy) ---
def _openai_headers():
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY missing")
    return {"Authorization": f"Bearer {OPENAI_API_KEY}"}

def _firestore():
    return firestore.Client(project=PROJECT_ID)

def _storage():
    return storage.Client(project=PROJECT_ID)

# --- Schemas (v1) ---
class IdeaReq(BaseModel):
    genre: str
    tone: str
    seed: str
    language: str = "en"

class OutlineReq(BaseModel):
    logline: str
    structure: str = "film"
    style: str = "Bollywood high-concept thriller"
    language: str = "en"

class ScriptReq(BaseModel):
    outline: str
    style: str
    language: str = "en"

class DeckReq(BaseModel):
    title: str
    logline: str
    synopsis: str
    characters: str
    world: str
    comps: str
    toneboard: Optional[str] = None
    language: str = "en"

class ExportReq(BaseModel):
    deck_json: Dict[str, Any]
    format: str  # pdf | docx

# --- Helpers ---
async def call_llm(system_prompt: str, user_prompt: str) -> str:
    """
    Robust wrapper around OpenAI Chat Completions.
    """
    url = "https://api.openai.com/v1/chat/completions"
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    try:
        async with httpx.AsyncClient(timeout=120) as client:
            r = await client.post(url, headers=_openai_headers(), json=payload)
            r.raise_for_status()
            data = r.json()
            content = data["choices"][0]["message"]["content"]
            if not isinstance(content, str):
                content = json.dumps(content, ensure_ascii=False)
            return content
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=502, detail=f"LLM upstream error: {e.response.text[:500]}")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"LLM request failed: {str(e)[:500]}")

def save_pdf_from_text(title: str, text: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.setTitle(title)
    y = 800
    for line in text.splitlines():
        if y < 50:
            c.showPage()
            y = 800
        c.drawString(40, y, line[:120])
        y -= 18
    c.showPage()
    c.save()
    return buf.getvalue()

def save_docx_from_text(title: str, text: str) -> bytes:
    doc = Document()
    doc.add_heading(title, 0)
    for line in text.splitlines():
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def _safe_path_segment(seg: str) -> str:
    # cheap sanitization for filenames/paths stored in GCS
    bad = ['..', '/', '\\', ':', '*', '?', '"', '<', '>', '|']
    for b in bad:
        seg = seg.replace(b, "_")
    return seg.strip() or "untitled"

def upload_bytes(path: str, data: bytes, content_type: str) -> str:
    client = _storage()
    bucket = client.bucket(BUCKET_NAME)
    blob = bucket.blob(path)
    blob.upload_from_string(data, content_type=content_type)
    blob.make_public()  # switch to signed URLs later if needed
    return blob.public_url

# -------- text extraction for uploads (pdf, docx, txt) ----------
async def _extract_text_from_upload(file: UploadFile) -> str:
    name = (file.filename or "upload").lower()
    raw = await file.read()
    ext = os.path.splitext(name)[1]
    def try_decode(b: bytes) -> str:
        return b.decode("utf-8", errors="ignore")

    if ext in [".txt", ".md"]:
        return try_decode(raw)

    if ext == ".pdf":
        with io.BytesIO(raw) as bio:
            reader = PdfReader(bio)
            return "\n".join((p.extract_text() or "") for p in reader.pages)

    if ext == ".docx":
        with io.BytesIO(raw) as bio:
            doc = Document(bio)
            return "\n".join(par.text for par in doc.paragraphs)

    # fallback
    return try_decode(raw)

# -------------- strict JSON helpers for v2 endpoints --------------
def _json_instructions(schema_desc: str) -> str:
    return (
        "Return ONLY valid JSON. No markdown fences, no commentary. "
        f"Schema: {schema_desc}"
    )

def _parse_json_fallback(s: str, default):
    try:
        return json.loads(s)
    except Exception:
        return default

# ===================== ROUTES =====================

@app.get("/")
def root():
    return {"status": "up", "service": "manthan-backend", "version": "0.3.0"}

@app.get("/health")
def health():
    return {
        "status": "ok",
        "version": "0.3.0",
        "project": PROJECT_ID,
        "bucket": BUCKET_NAME,
        "namespace": NAMESPACE,
    }

@app.get("/runtime-env")
def runtime_env():
    """Return only the safe Firebase env vars to the frontend."""
    return {
        "FIREBASE_API_KEY": os.getenv("NEXT_PUBLIC_FIREBASE_API_KEY", ""),
        "FIREBASE_AUTH_DOMAIN": os.getenv("NEXT_PUBLIC_FIREBASE_AUTH_DOMAIN", ""),
        "FIREBASE_PROJECT_ID": os.getenv("NEXT_PUBLIC_FIREBASE_PROJECT_ID", ""),
        "FIREBASE_APP_ID": os.getenv("NEXT_PUBLIC_FIREBASE_APP_ID", ""),
        "FIREBASE_STORAGE_BUCKET": os.getenv("NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET", ""),
        "FIREBASE_MESSAGING_SENDER_ID": os.getenv("NEXT_PUBLIC_FIREBASE_MESSAGING_SENDER_ID", ""),
    }

# -------- v1 generation endpoints (kept for backward-compat) --------
@app.post("/gen/idea")
async def gen_idea(req: IdeaReq):
    sys = "You are a multilingual film/series ideation assistant for Indian markets. Return 5 sharp loglines and a 1-paragraph premise."
    user = f"Language: {req.language}\nGenre: {req.genre}\nTone: {req.tone}\nSeed: {req.seed}"
    content = await call_llm(sys, user)
    return {"content": content}

@app.post("/gen/outline")
async def gen_outline(req: OutlineReq):
    sys = "You are a professional story editor. Produce a beat-sheet with acts/episodes, character bios, and world notes. Keep it concise and production-ready."
    user = f"Structure: {req.structure}\nStyle: {req.style}\nLanguage: {req.language}\nLogline: {req.logline}"
    content = await call_llm(sys, user)
    return {"content": content}

@app.post("/gen/script")
async def gen_script(req: ScriptReq):
    sys = "You are a screenwriter. Expand the outline into screenplay pages in Fountain-like format. Keep dialogue punchy; format with scene headers."
    user = f"Style: {req.style}\nLanguage: {req.language}\nOutline:\n{req.outline}"
    content = await call_llm(sys, user)
    return {"content": content}

@app.post("/gen/deck")
async def gen_deck(req: DeckReq):
    sys = "You are a pitch-deck producer for film/series. Create structured JSON with slides: Title, Logline, Synopsis, Characters, World, Toneboard, Comparables, CTA."
    user = json.dumps(req.model_dump(), ensure_ascii=False)
    content = await call_llm(sys, user)
    try:
        deck = json.loads(content)
    except Exception:
        deck = {"raw": content}
    return {"deck": deck}

# --------------- improved ingest (pdf/docx/txt) ----------------
@app.post("/ingest/upload")
async def ingest_upload(file: UploadFile = File(...), language: str = Form("en")):
    text = await _extract_text_from_upload(file)
    sys = ("You are a development exec. From the uploaded script text, extract: "
           "title, logline, synopsis, characters, world, themes, and comparables.")
    user = f"Language: {language}\nTEXT:\n{text[:15000]}"
    content = await call_llm(sys, user)
    return {"extracted": content, "filename": file.filename}

# ---------------------- exports ----------------------
@app.post("/export")
async def export(req: ExportReq):
    deck = req.deck_json
    title = deck.get("title") or (deck.get("raw", "Pitch Deck")[:40])
    safe_title = _safe_path_segment(title)

    text = json.dumps(deck, ensure_ascii=False, indent=2) if "raw" not in deck else str(deck["raw"])

    if req.format == "pdf":
        data = save_pdf_from_text(safe_title, text)
        url = upload_bytes(f"exports/{safe_title}.pdf", data, "application/pdf")
        return {"url": url}
    elif req.format == "docx":
        data = save_docx_from_text(safe_title, text)
        url = upload_bytes(
            f"exports/{safe_title}.docx", data,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return {"url": url}
    else:
        raise HTTPException(status_code=400, detail="unsupported format")

# ---------------------- image stub ----------------------
@app.post("/images/concepts")
async def concepts(prompt: str = Form(...)):
    if not ENABLE_IMAGE_GEN:
        return {"enabled": False}
    return {"enabled": True, "images": [f"https://storage.googleapis.com/{BUCKET_NAME}/concepts/sample1.png"]}

# ===================== v2 JSON FLOW =====================

# ---- v2 Ideas (3 options) ----
class IdeasV2Req(BaseModel):
    genre: str
    tone: str
    seed: str
    language: str = "en"

@app.post("/v2/gen/ideas")
async def v2_gen_ideas(req: IdeasV2Req):
    sys = (
        "You are a multilingual film/series ideation assistant for Indian markets. " +
        _json_instructions(
            '{ "options": [ {"logline": "string", "premise": "string"} ] } '
            'with exactly 3 items in "options"'
        )
    )
    user = (
        f"Language: {req.language}\nGenre: {req.genre}\nTone: {req.tone}\nSeed: {req.seed}\n"
        "Generate 3 distinct ideas consistent with Indian audience sensibilities."
    )
    content = await call_llm(sys, user)
    data = _parse_json_fallback(content, {"options": []})
    data["options"] = (data.get("options") or [])[:3]
    return data

# ---- v2 Outlines (3 options) ----
class OutlineV2Req(BaseModel):
    logline: str
    structure: str = "film"
    style: str = "Bollywood high-concept thriller"
    language: str = "en"

@app.post("/v2/gen/outlines")
async def v2_gen_outlines(req: OutlineV2Req):
    sys = (
        "You are a professional story editor. " +
        _json_instructions(
            '{ "options": [ {"outline": "string"} ] } with exactly 3 items'
        )
    )
    user = (
        f"Structure: {req.structure}\nStyle: {req.style}\nLanguage: {req.language}\n"
        f"Logline:\n{req.logline}\n"
        "Create crisp beats with acts, key turns, character beats, and world notes."
    )
    content = await call_llm(sys, user)
    data = _parse_json_fallback(content, {"options": []})
    data["options"] = (data.get("options") or [])[:3]
    return data

# ---- v2 Scripts (3 options) ----
class ScriptV2Req(BaseModel):
    outline: str
    style: str
    language: str = "en"

@app.post("/v2/gen/scripts")
async def v2_gen_scripts(req: ScriptV2Req):
    sys = (
        "You are a screenwriter. " +
        _json_instructions(
            '{ "options": [ {"script": "string"} ] } '
            "with exactly 3 items (opening 3–5 pages in Fountain-like format)"
        )
    )
    user = (
        f"Style: {req.style}\nLanguage: {req.language}\n"
        f"Outline:\n{req.outline}\n"
        "Write the opening 3-5 pages. Punchy dialogue. Scene headers."
    )
    content = await call_llm(sys, user)
    data = _parse_json_fallback(content, {"options": []})
    data["options"] = (data.get("options") or [])[:3]
    return data

# ---- v2 Deck build (1 option) ----
class DeckBuildV2Req(BaseModel):
    title: str
    logline: str
    synopsis: str
    characters: str
    world: str
    comps: str
    toneboard: Optional[str] = None
    language: str = "en"

@app.post("/v2/deck/build")
async def v2_deck_build(req: DeckBuildV2Req):
    sys = (
        "You are a pitch-deck producer for film/series. " +
        _json_instructions(
            '{ "options": [ {"deck": {'
            '"title":"string","logline":"string","synopsis":"string","characters":"string",'
            '"world":"string","comps":"string","toneboard":"string","cta":"string"}} ] } '
            "with exactly 1 item"
        )
    )
    user = json.dumps(req.model_dump(), ensure_ascii=False)
    content = await call_llm(sys, user)
    data = _parse_json_fallback(content, {"options": [{"deck": req.model_dump()}]})
    data["options"] = (data.get("options") or [])[:1]
    return data

# ===================== Firestore: Projects API =====================

# Data will be saved under:
#   <NAMESPACE>_users/{uid}/projects/{project_id}
# Steps can be stored in a subcollection "steps" or merged as fields.

USERS_COLL = f"{NAMESPACE}_users"

class CreateProjectReq(BaseModel):
    uid: str
    title: Optional[str] = None

@app.post("/v2/projects/create")
def create_project(req: CreateProjectReq):
    db = _firestore()
    project_id = uuid.uuid4().hex[:12]
    doc_ref = db.collection(USERS_COLL).document(req.uid).collection("projects").document(project_id)
    doc_ref.set({
        "project_id": project_id,
        "title": req.title or "Untitled Project",
        "created_at": firestore.SERVER_TIMESTAMP,
        "updated_at": firestore.SERVER_TIMESTAMP,
    })
    return {"project_id": project_id}

@app.get("/v2/projects/list")
def list_projects(uid: str = Query(..., description="Firebase UID")):
    db = _firestore()
    col = db.collection(USERS_COLL).document(uid).collection("projects")
    docs = col.order_by("updated_at", direction=firestore.Query.DESCENDING).limit(50).stream()
    items = []
    for d in docs:
        data = d.to_dict() or {}
        data["id"] = d.id
        items.append(data)
    return {"projects": items}

class SaveStepReq(BaseModel):
    uid: str
    project_id: str
    step: str                   # "ideas" | "outline" | "script" | "deck"
    payload: Dict[str, Any]

@app.post("/v2/projects/save-step")
def save_step(req: SaveStepReq):
    db = _firestore()
    base = db.collection(USERS_COLL).document(req.uid).collection("projects").document(req.project_id)
    # store step payload under a nested map and keep updated_at fresh
    base.set({
        "steps": { req.step: req.payload },
        "updated_at": firestore.SERVER_TIMESTAMP,
    }, merge=True)
    return {"ok": True}

@app.get("/v2/projects/get")
def get_project(uid: str = Query(...), project_id: str = Query(...)):
    db = _firestore()
    doc = db.collection(USERS_COLL).document(uid).collection("projects").document(project_id).get()
    if not doc.exists:
        raise HTTPException(status_code=404, detail="project not found")
    data = doc.to_dict() or {}
    data["id"] = doc.id
    return data



