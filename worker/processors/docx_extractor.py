def extract(file_bytes: bytes) -> str:
    from docx import Document
    import io

    try:
        doc = Document(io.BytesIO(file_bytes))
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style and para.style.name else ""
                if style_name.startswith("Heading"):
                    lines.append(f"\n## {text}\n")
                else:
                    lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)

        result = "\n".join(lines).strip()
        if not result:
            raise ValueError("No readable text found in DOCX")
        return result
    except Exception as exc:
        raise Exception(f"DOCX extraction failed: {exc}") from exc
